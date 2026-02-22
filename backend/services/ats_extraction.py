"""
ATS CV Text Extraction Service
Extracts text from PDF, DOC, DOCX files for ATS processing
"""

import os
import re
from typing import Optional, Tuple

# Minimum characters for a CV to be considered readable
MIN_TEXT_LENGTH = 200


async def extract_text_from_file(file_path: str) -> Tuple[str, bool, str]:
    """
    Extract text from CV file.
    
    Returns:
        Tuple of (extracted_text, is_readable, error_message)
    """
    if not os.path.exists(file_path):
        return "", False, "File not found"
    
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == '.pdf':
            text = await extract_from_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            text = await extract_from_docx(file_path)
        else:
            return "", False, f"Unsupported file type: {ext}"
        
        # Clean text
        text = clean_text(text)
        
        # Check if readable
        if len(text) < MIN_TEXT_LENGTH:
            return text, False, "Text too short - likely scanned/image PDF"
        
        return text, True, ""
        
    except Exception as e:
        return "", False, str(e)


async def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF using multiple methods"""
    text = ""
    
    # Try pdfplumber first (best for text PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if len(text.strip()) > MIN_TEXT_LENGTH:
            return text
    except Exception:
        pass
    
    # Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception:
        pass
    
    return text


async def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX/DOC files"""
    text = ""
    
    try:
        from docx import Document
        doc = Document(file_path)
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
    except Exception as e:
        # For .doc files, try alternative
        if file_path.endswith('.doc'):
            try:
                import subprocess
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                text = result.stdout
            except Exception:
                pass
    
    return text


def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep Arabic
    text = re.sub(r'[^\w\s\u0600-\u06FF\u0750-\u077F.,;:!?@#$%&*()-]', '', text)
    
    return text.strip()


def detect_language(text: str) -> str:
    """Detect primary language of text (ar/en/mixed)"""
    if not text:
        return "unknown"
    
    arabic_chars = len(re.findall(r'[\u0600-\u06FF]', text))
    english_chars = len(re.findall(r'[a-zA-Z]', text))
    
    total = arabic_chars + english_chars
    if total == 0:
        return "unknown"
    
    arabic_ratio = arabic_chars / total
    
    if arabic_ratio > 0.7:
        return "ar"
    elif arabic_ratio < 0.3:
        return "en"
    else:
        return "mixed"
